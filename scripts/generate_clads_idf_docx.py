from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_TXT = ROOT / "IDF_B_CLADS.txt"
OUTPUT_DOC = ROOT / "output" / "doc" / "CLADS_Invention_Disclosure_Format_B_v2.docx"


def normalize_text(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2011": "-",
        ">= ": ">= ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    for idx in range(1, 10):
        text = text.replace(f"beta_{idx}", f"\u03b2{to_subscript(str(idx))}")
    text = text.replace(">=", "\u2265")
    text = text.replace("<=", "\u2264")
    return text


def to_subscript(value: str) -> str:
    table = str.maketrans("0123456789+-=()aehijklmnoprstuvx", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")
    return value.translate(table)


def set_page_layout(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def configure_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for section in doc.sections:
        set_page_layout(section)


def style_run(run, *, bold: bool = False, italic: bool = False, size: int | None = None, font: str | None = None) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = font or "Times New Roman"
    run.font.size = Pt(size or 11)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "4F81BD")


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VIT IPR&TT CELL")
    style_run(r, bold=True, size=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Invention Disclosure Format (IDF)-B")
    style_run(r, bold=True, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document No. 02-IPR-R003 | Issue No/Date: 2 / 01.02.2024 | Amd. No/Date: 0 / 00.00.0000")
    style_run(r, size=9)


def parse_ascii_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in lines:
        stripped = raw.strip()
        if not stripped.startswith("|"):
            continue
        cells = [normalize_text(cell.strip()) for cell in stripped.strip("|").split("|")]
        if not any(cells):
            continue
        if rows and not cells[0]:
            previous = rows[-1]
            for idx, value in enumerate(cells):
                if value:
                    previous[idx] = f"{previous[idx]} {value}".strip()
        else:
            rows.append(cells)
    return rows


def write_table(doc: Document, rows: list[list[str]], *, font_size: int = 9, header_fill: str = "D9EAF7") -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if row_idx == 0:
                shade_cell(cell, header_fill)
            p = cell.paragraphs[0]
            if row_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_plain_text(p, value, bold=(row_idx == 0), size=font_size)
    doc.add_paragraph()


def add_plain_text(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    run = paragraph.add_run(normalize_text(text))
    style_run(run, bold=bold, italic=italic, size=size)


MATH_PATTERN = re.compile(r"beta_(\d+)|([A-Za-z]+)_([A-Za-z0-9]+)")


def add_math_text(paragraph, text: str, *, size: int = 12) -> None:
    text = text.replace(">=", "\u2265").replace("<=", "\u2264")
    cursor = 0
    for match in MATH_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            style_run(run, size=size, font="Cambria Math")
        if match.group(1):
            base = paragraph.add_run("\u03b2")
            style_run(base, italic=True, size=size, font="Cambria Math")
            sub = paragraph.add_run(match.group(1))
            style_run(sub, size=size - 1, font="Cambria Math")
            sub.font.subscript = True
        else:
            base = paragraph.add_run(match.group(2))
            style_run(base, italic=True, size=size, font="Cambria Math")
            sub = paragraph.add_run(match.group(3))
            style_run(sub, size=size - 1, font="Cambria Math")
            sub.font.subscript = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        style_run(run, size=size, font="Cambria Math")


def add_equation_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_math_text(p, line.strip(), size=12)
    doc.add_paragraph()


FIGURE_RE = re.compile(r"^\[Figure\s+(\d+):\s*(.+?)\s*-\s*see\s*(.+?)\]$")


def add_figure(doc: Document, line: str) -> None:
    normalized = normalize_text(line.strip())
    match = FIGURE_RE.match(normalized)
    if not match:
        add_body_paragraph(doc, normalized)
        return
    _, caption, rel_path = match.groups()
    image_path = ROOT / Path(rel_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        p.add_run().add_picture(str(image_path), width=Inches(6.75))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_plain_text(cap, f"Figure {match.group(1)}. {caption}", italic=True, size=10)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    size = 13 if level == 1 else 12
    run = p.add_run(normalize_text(text))
    style_run(run, bold=True, size=size)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if ":" in text and re.match(r"^(Summary|Background|Results Summary|Current Status|Experiment \d+ - .+|Claim \d+ .+):", text):
        label, remainder = text.split(":", 1)
        add_plain_text(p, label + ":", bold=True)
        if remainder.strip():
            add_plain_text(p, " " + remainder.strip())
    else:
        add_plain_text(p, text)


def add_numbered_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(4)
    add_plain_text(p, text)


def is_heading(line: str) -> bool:
    return bool(re.match(r"^\d+\.\s", line.strip()) and line.strip().endswith(":")) or bool(re.match(r"^\d+\.\d+\s", line.strip()))


def is_numbered_item(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r"^\d+\.\s", stripped)) and not stripped.endswith(":") and not re.match(r"^\d+\.\d+\s", stripped)


def collect_paragraph(lines: list[str], start: int) -> tuple[str, int]:
    parts: list[str] = []
    idx = start
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            break
        if stripped.startswith("[Figure "):
            break
        if stripped.startswith("+") or stripped.startswith("|"):
            break
        if is_heading(stripped):
            break
        if idx != start and (line.startswith("    ") and "=" in stripped):
            break
        if idx != start and is_numbered_item(stripped):
            break
        parts.append(stripped)
        idx += 1
    return normalize_text(" ".join(parts)), idx


def collect_numbered_item(lines: list[str], start: int) -> tuple[str, int]:
    first = lines[start].strip()
    parts = [first]
    idx = start + 1
    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            break
        if is_heading(stripped) or is_numbered_item(stripped) or stripped.startswith("[Figure "):
            break
        if stripped.startswith("+") or stripped.startswith("|"):
            break
        parts.append(stripped)
        idx += 1
    return normalize_text(" ".join(parts)), idx


def collect_equation(lines: list[str], start: int) -> tuple[list[str], int]:
    block: list[str] = []
    idx = start
    while idx < len(lines):
        line = lines[idx]
        if not line.startswith("    "):
            break
        stripped = line.strip()
        if not stripped:
            break
        block.append(normalize_text(stripped))
        idx += 1
    return block, idx


def add_section_break(doc: Document, *, landscape: bool = False) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_layout(section, landscape=landscape)


def build_doc() -> Path:
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)

    text = SOURCE_TXT.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    metadata_block: list[str] = []
    idx = 0
    while idx < len(lines) and "Inventor(s) Name(s)" not in lines[idx]:
        idx += 1
    while idx < len(lines) and lines[idx].strip():
        if lines[idx].strip().startswith(("+", "|")):
            metadata_block.append(lines[idx])
        idx += 1
    write_table(doc, parse_ascii_table(metadata_block), font_size=10)

    while idx < len(lines) and not lines[idx].strip().startswith("1. Title"):
        idx += 1

    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped or stripped.startswith("===") or "END OF THE DOCUMENT" in stripped:
            idx += 1
            continue
        if stripped.startswith("+") or stripped.startswith("|"):
            block: list[str] = []
            while idx < len(lines) and lines[idx].strip().startswith(("+", "|")):
                block.append(lines[idx])
                idx += 1
            rows = parse_ascii_table(block)
            wide = len(rows[0]) >= 5 if rows else False
            if wide:
                add_section_break(doc, landscape=True)
            write_table(doc, rows, font_size=8 if wide else 9)
            if wide:
                add_section_break(doc, landscape=False)
            continue
        if stripped.startswith("[Figure "):
            add_figure(doc, stripped)
            idx += 1
            continue
        if is_heading(stripped):
            add_heading(doc, stripped.rstrip(":"), 1 if stripped.count(".") == 1 else 2)
            idx += 1
            continue
        if lines[idx].startswith("    ") and "=" in stripped:
            block, idx = collect_equation(lines, idx)
            add_equation_block(doc, block)
            continue
        if is_numbered_item(stripped):
            item, idx = collect_numbered_item(lines, idx)
            add_numbered_item(doc, item)
            continue
        paragraph, idx = collect_paragraph(lines, idx)
        if paragraph:
            add_body_paragraph(doc, paragraph)
        else:
            idx += 1

    doc.save(OUTPUT_DOC)
    return OUTPUT_DOC


if __name__ == "__main__":
    path = build_doc()
    print(path)
